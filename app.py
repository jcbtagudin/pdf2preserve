from flask import Flask, request, send_file, jsonify, session
import fitz  # PyMuPDF
import os
import re
import time
import zipfile
import tempfile
import json
import uuid
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'pdf2preserve_secret_key_change_in_production')
UPLOAD_FOLDER = "/tmp/uploads" if os.environ.get("RAILWAY_ENVIRONMENT") else "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Export tracking storage (in production, use a database)
export_tracking = {}

# Export limits
GUEST_DAILY_LIMIT = 10
LOGGED_IN_DAILY_LIMIT = 30

def get_user_key():
    """Get unique identifier for tracking exports"""
    if 'user_id' in session:
        return f"user_{session['user_id']}"
    else:
        # For guests, use session ID
        if 'session_id' not in session:
            session['session_id'] = str(uuid.uuid4())
        return f"guest_{session['session_id']}"

def is_logged_in():
    """Check if user is logged in"""
    return 'user_id' in session

def get_daily_limit():
    """Get daily export limit based on user status"""
    return LOGGED_IN_DAILY_LIMIT if is_logged_in() else GUEST_DAILY_LIMIT

def check_and_update_exports(user_key):
    """Check if user can export and update count"""
    now = datetime.now()
    today = now.date()
    
    if user_key not in export_tracking:
        export_tracking[user_key] = {
            'count': 0,
            'date': today.isoformat(),
            'last_export': now.isoformat()
        }
    
    user_data = export_tracking[user_key]
    last_date = datetime.fromisoformat(user_data['date']).date()
    
    # Reset if new day
    if today > last_date:
        user_data['count'] = 0
        user_data['date'] = today.isoformat()
    
    limit = get_daily_limit()
    
    if user_data['count'] >= limit:
        return False, user_data['count'], limit
    
    # Increment count
    user_data['count'] += 1
    user_data['last_export'] = now.isoformat()
    
    return True, user_data['count'], limit

def get_export_status(user_key):
    """Get current export status without incrementing"""
    today = datetime.now().date()
    
    if user_key not in export_tracking:
        return 0, get_daily_limit()
    
    user_data = export_tracking[user_key]
    last_date = datetime.fromisoformat(user_data['date']).date()
    
    # Reset if new day
    if today > last_date:
        return 0, get_daily_limit()
    
    return user_data['count'], get_daily_limit()

class PDFFormatter:
    def __init__(self, pdf_path, layout_mode='preserve'):
        self.doc = fitz.open(pdf_path)
        self.formatted_content = []
        self.layout_mode = layout_mode
        self.bullet_patterns = ['•', '●', '◦', '▪', '▫', '■', '□', '◆', '◇', '-', '*']
        
    def extract_with_formatting(self):
        all_blocks = []
        
        for page_num in range(len(self.doc)):
            page = self.doc[page_num]
            page_rect = page.rect
            page_width = page_rect.width
            blocks = page.get_text("dict")["blocks"]
            
            # Extract tables first if in preserve mode
            tables = []
            if self.layout_mode == 'preserve':
                try:
                    page_tables = page.find_tables()
                    for table in page_tables:
                        table_data = table.extract()
                        table_bbox = table.bbox
                        tables.append({
                            'data': table_data,
                            'bbox': table_bbox,
                            'type': 'table'
                        })
                except:
                    pass  # Table detection may fail on some PDFs
            
            for block in blocks:
                if "lines" in block:
                    block_bbox = block.get("bbox", [0, 0, 0, 0])
                    
                    # Skip blocks that are part of detected tables
                    is_table_block = False
                    for table in tables:
                        if self._bbox_overlap(block_bbox, table['bbox']):
                            is_table_block = True
                            break
                    
                    if is_table_block:
                        continue
                    
                    block_content = []
                    block_y_positions = []
                    
                    for line in block["lines"]:
                        line_content = []
                        line_bbox = line.get("bbox", [0, 0, 0, 0])
                        block_y_positions.append(line_bbox[1])  # y0 coordinate
                        
                        # Detect text alignment for this line
                        line_alignment = self._detect_text_alignment(line_bbox, page_width, block_bbox)
                        
                        for span in line["spans"]:
                            text = span["text"].strip()
                            if text:
                                flags = span["flags"]
                                size = span["size"]
                                font = span.get("font", "")
                                
                                is_bold = bool(flags & 2**4)
                                is_italic = bool(flags & 2**1)
                                
                                # Enhanced heading detection
                                avg_size = 12  # Assumed average text size
                                is_heading = (size > avg_size + 2) or (is_bold and size >= avg_size)
                                heading_level = self._calculate_heading_level(size, is_bold, font)
                                
                                # Detect list items
                                is_list_item = self._is_list_item(text)
                                list_type = self._get_list_type(text) if is_list_item else None
                                
                                line_content.append({
                                    'text': text,
                                    'bold': is_bold,
                                    'italic': is_italic,
                                    'size': size,
                                    'font': font,
                                    'heading': heading_level,
                                    'is_list': is_list_item,
                                    'list_type': list_type,
                                    'alignment': line_alignment,
                                    'bbox': line_bbox
                                })
                        
                        if line_content:
                            block_content.append(line_content)
                    
                    if block_content:
                        # Detect paragraph breaks based on spacing
                        block_type = self._detect_block_type(block_content, block_y_positions)
                        all_blocks.append({
                            'content': block_content,
                            'type': block_type,
                            'bbox': block_bbox
                        })
            
            # Add tables to blocks
            for table in tables:
                all_blocks.append(table)
        
        self.formatted_content = all_blocks
        self.doc.close()
        return self.formatted_content
    
    def _bbox_overlap(self, bbox1, bbox2):
        """Check if two bounding boxes overlap"""
        return not (bbox1[2] < bbox2[0] or bbox2[2] < bbox1[0] or 
                   bbox1[3] < bbox2[1] or bbox2[3] < bbox1[1])
    
    def _calculate_heading_level(self, size, is_bold, font):
        """Calculate heading level based on font properties"""
        if size <= 10:
            return 0
        elif size >= 20:
            return 1
        elif size >= 18:
            return 2
        elif size >= 16:
            return 3
        elif size >= 14 or (is_bold and size >= 12):
            return 4
        elif is_bold and size >= 11:
            return 5
        elif size >= 13:
            return 6
        return 0
    
    def _is_list_item(self, text):
        """Detect if text starts with bullet point or number"""
        if not text:
            return False
        
        # Check for bullet points
        for bullet in self.bullet_patterns:
            if text.startswith(bullet + ' ') or text.startswith(bullet + '\t'):
                return True
        
        # Check for numbered lists (1., 1), a., a), i., etc.)
        import re
        numbered_patterns = [
            r'^\d+\.\s',  # 1. 2. 3.
            r'^\d+\)\s',  # 1) 2) 3)
            r'^[a-zA-Z]\.\s',  # a. b. c.
            r'^[a-zA-Z]\)\s',  # a) b) c)
            r'^[ivxlcdm]+\.\s',  # i. ii. iii. (roman numerals)
            r'^[IVXLCDM]+\.\s'   # I. II. III.
        ]
        
        for pattern in numbered_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _get_list_type(self, text):
        """Determine the type of list (bullet, numbered, etc.)"""
        if not text:
            return None
        
        for bullet in self.bullet_patterns:
            if text.startswith(bullet):
                return 'bullet'
        
        import re
        if re.match(r'^\d+[\.\)]\s', text):
            return 'numbered'
        elif re.match(r'^[a-zA-Z][\.\)]\s', text):
            return 'lettered'
        elif re.match(r'^[ivxlcdmIVXLCDM]+\.\s', text):
            return 'roman'
        
        return 'bullet'
    
    def _detect_text_alignment(self, line_bbox, page_width, block_bbox=None):
        """Detect text alignment based on position within page/block"""
        if not line_bbox or len(line_bbox) < 4:
            return 'left'
        
        x0, y0, x1, y1 = line_bbox
        line_width = x1 - x0
        
        # Use block width if available, otherwise use page width
        container_width = page_width
        container_left = 0
        
        if block_bbox and len(block_bbox) >= 4:
            container_left = block_bbox[0]
            container_width = block_bbox[2] - block_bbox[0]
        
        # Calculate relative position within container
        line_left_margin = x0 - container_left
        line_right_margin = (container_left + container_width) - x1
        
        # Define thresholds for alignment detection
        center_threshold = 0.1 * container_width  # 10% of container width
        margin_threshold = 0.05 * container_width  # 5% of container width
        
        # Detect alignment based on margins
        if abs(line_left_margin - line_right_margin) < center_threshold:
            # Text is roughly centered
            return 'center'
        elif line_right_margin < margin_threshold and line_left_margin > container_width * 0.3:
            # Text is close to right edge with significant left margin
            return 'right'
        elif line_left_margin < margin_threshold:
            # Text is close to left edge
            return 'left'
        else:
            # Default to left if unclear
            return 'left'
    
    def _detect_block_type(self, block_content, y_positions):
        """Detect if block is paragraph, heading, list, etc."""
        if not block_content:
            return 'text'
        
        first_line = block_content[0]
        if first_line and first_line[0].get('heading', 0) > 0:
            return 'heading'
        
        if any(line and line[0].get('is_list', False) for line in block_content):
            return 'list'
        
        return 'paragraph'
    
    def to_html(self):
        html_content = ['<!DOCTYPE html>', '<html>', '<head>', '<meta charset="UTF-8">', 
                       '<title>PDF Content</title>', 
                       '<style>table { border-collapse: collapse; width: 100%; margin: 20px 0; } th, td { border: 1px solid #ddd; padding: 8px; text-align: left; } th { background-color: #f2f2f2; }</style>',
                       '</head>', '<body>']
        
        if self.layout_mode == 'clean':
            return self._to_html_clean()
        
        for block in self.formatted_content:
            if block.get('type') == 'table':
                html_content.append(self._table_to_html(block['data']))
            else:
                block_content = block.get('content', [])
                block_type = block.get('type', 'paragraph')
                
                if block_type == 'heading':
                    html_content.extend(self._block_to_html_heading(block_content))
                elif block_type == 'list':
                    html_content.extend(self._block_to_html_list(block_content))
                else:
                    html_content.extend(self._block_to_html_paragraph(block_content))
        
        html_content.extend(['</body>', '</html>'])
        return '\n'.join(html_content)
    
    def _to_html_clean(self):
        """Simple clean HTML without structure"""
        html_content = ['<!DOCTYPE html>', '<html>', '<head>', '<meta charset="UTF-8">', 
                       '<title>PDF Content</title>', '</head>', '<body>', '<div>']
        
        all_text = []
        for block in self.formatted_content:
            if block.get('type') == 'table':
                continue  # Skip tables in clean mode
            
            block_content = block.get('content', [])
            for line in block_content:
                line_text = ' '.join([span['text'] for span in line])
                if line_text.strip():
                    all_text.append(line_text.strip())
        
        html_content.append('<p>' + ' '.join(all_text) + '</p>')
        html_content.extend(['</div>', '</body>', '</html>'])
        return '\n'.join(html_content)
    
    def _block_to_html_heading(self, block_content):
        if not block_content or not block_content[0]:
            return []
        
        heading_level = block_content[0][0].get('heading', 3)
        alignment = block_content[0][0].get('alignment', 'left')
        text_content = ' '.join([' '.join([span['text'] for span in line]) for line in block_content])
        
        align_style = f' style="text-align: {alignment};"' if alignment != 'left' else ''
        return [f'<h{heading_level}{align_style}>{text_content}</h{heading_level}>']
    
    def _block_to_html_list(self, block_content):
        html_lines = []
        current_list_type = None
        list_open = False
        
        for line in block_content:
            if not line:
                continue
                
            first_span = line[0]
            if first_span.get('is_list'):
                list_type = first_span.get('list_type', 'bullet')
                
                # Start new list if needed
                if not list_open or list_type != current_list_type:
                    if list_open:
                        html_lines.append('</ul>' if current_list_type == 'bullet' else '</ol>')
                    
                    if list_type == 'bullet':
                        html_lines.append('<ul>')
                    else:
                        html_lines.append('<ol>')
                    
                    list_open = True
                    current_list_type = list_type
                
                # Remove bullet/number from text
                text_content = first_span['text']
                for bullet in self.bullet_patterns:
                    if text_content.startswith(bullet):
                        text_content = text_content[len(bullet):].strip()
                        break
                
                import re
                text_content = re.sub(r'^[\d\w]+[\.\)]\s*', '', text_content)
                
                # Add remaining spans
                for span in line[1:]:
                    text_content += ' ' + span['text']
                
                html_lines.append(f'<li>{text_content.strip()}</li>')
            else:
                # Regular text in list context
                text_content = ' '.join([span['text'] for span in line])
                if text_content.strip():
                    html_lines.append(f'<p>{text_content}</p>')
        
        if list_open:
            html_lines.append('</ul>' if current_list_type == 'bullet' else '</ol>')
        
        return html_lines
    
    def _block_to_html_paragraph(self, block_content):
        if not block_content:
            return []
        
        # Get alignment from first line
        alignment = block_content[0][0].get('alignment', 'left') if block_content[0] else 'left'
        align_style = f' style="text-align: {alignment};"' if alignment != 'left' else ''
        
        p_content = f'<p{align_style}>'
        for line in block_content:
            for span in line:
                text = span['text']
                if span['bold'] and span['italic']:
                    text = f'<strong><em>{text}</em></strong>'
                elif span['bold']:
                    text = f'<strong>{text}</strong>'
                elif span['italic']:
                    text = f'<em>{text}</em>'
                p_content += text + ' '
            p_content += '<br>' if len(block_content) > 1 else ''
        
        p_content += '</p>'
        return [p_content]
    
    def _table_to_html(self, table_data):
        if not table_data:
            return ''
        
        html = ['<table>']
        
        # First row as header if it looks like headers
        first_row = table_data[0] if table_data else []
        has_headers = len(first_row) > 0 and any(cell and isinstance(cell, str) and cell.strip() for cell in first_row)
        
        if has_headers:
            html.append('<thead><tr>')
            for cell in first_row:
                html.append(f'<th>{cell or ""}</th>')
            html.append('</tr></thead>')
            data_rows = table_data[1:]
        else:
            data_rows = table_data
        
        html.append('<tbody>')
        for row in data_rows:
            html.append('<tr>')
            for cell in row:
                html.append(f'<td>{cell or ""}</td>')
            html.append('</tr>')
        html.append('</tbody>')
        
        html.append('</table>')
        return '\n'.join(html)
    
    def to_markdown(self):
        if self.layout_mode == 'clean':
            return self._to_markdown_clean()
        
        markdown_content = []
        
        for block in self.formatted_content:
            if block.get('type') == 'table':
                markdown_content.append(self._table_to_markdown(block['data']))
            else:
                block_content = block.get('content', [])
                block_type = block.get('type', 'paragraph')
                
                if block_type == 'heading':
                    markdown_content.extend(self._block_to_markdown_heading(block_content))
                elif block_type == 'list':
                    markdown_content.extend(self._block_to_markdown_list(block_content))
                else:
                    markdown_content.extend(self._block_to_markdown_paragraph(block_content))
        
        return '\n\n'.join([item for item in markdown_content if item.strip()])
    
    def _to_markdown_clean(self):
        all_text = []
        for block in self.formatted_content:
            if block.get('type') == 'table':
                continue
            
            block_content = block.get('content', [])
            for line in block_content:
                line_text = ' '.join([span['text'] for span in line])
                if line_text.strip():
                    all_text.append(line_text.strip())
        
        return ' '.join(all_text)
    
    def _block_to_markdown_heading(self, block_content):
        if not block_content or not block_content[0]:
            return []
        
        heading_level = block_content[0][0].get('heading', 3)
        alignment = block_content[0][0].get('alignment', 'left')
        text_content = ' '.join([' '.join([span['text'] for span in line]) for line in block_content])
        
        heading_text = '#' * heading_level + ' ' + text_content
        
        # Add HTML alignment for non-left aligned headings in Markdown
        if alignment == 'center':
            heading_text += '\n<div align="center">' + text_content + '</div>'
            return [heading_text]
        elif alignment == 'right':
            heading_text += '\n<div align="right">' + text_content + '</div>'
            return [heading_text]
        
        return [heading_text]
    
    def _block_to_markdown_list(self, block_content):
        markdown_lines = []
        
        for line in block_content:
            if not line:
                continue
                
            first_span = line[0]
            if first_span.get('is_list'):
                list_type = first_span.get('list_type', 'bullet')
                
                # Remove bullet/number from text
                text_content = first_span['text']
                for bullet in self.bullet_patterns:
                    if text_content.startswith(bullet):
                        text_content = text_content[len(bullet):].strip()
                        break
                
                import re
                text_content = re.sub(r'^[\d\w]+[\.\)]\s*', '', text_content)
                
                # Add remaining spans with formatting
                for span in line[1:]:
                    span_text = span['text']
                    if span['bold'] and span['italic']:
                        span_text = f'***{span_text}***'
                    elif span['bold']:
                        span_text = f'**{span_text}**'
                    elif span['italic']:
                        span_text = f'*{span_text}*'
                    text_content += ' ' + span_text
                
                if list_type == 'numbered':
                    markdown_lines.append(f'1. {text_content.strip()}')
                else:
                    markdown_lines.append(f'- {text_content.strip()}')
            else:
                # Regular text
                line_text = ''
                for span in line:
                    text = span['text']
                    if span['bold'] and span['italic']:
                        text = f'***{text}***'
                    elif span['bold']:
                        text = f'**{text}**'
                    elif span['italic']:
                        text = f'*{text}*'
                    line_text += text + ' '
                
                if line_text.strip():
                    markdown_lines.append(line_text.strip())
        
        return markdown_lines
    
    def _block_to_markdown_paragraph(self, block_content):
        if not block_content:
            return []
        
        # Get alignment from first line
        alignment = block_content[0][0].get('alignment', 'left') if block_content[0] else 'left'
        
        paragraph_lines = []
        for line in block_content:
            line_text = ''
            for span in line:
                text = span['text']
                if span['bold'] and span['italic']:
                    text = f'***{text}***'
                elif span['bold']:
                    text = f'**{text}**'
                elif span['italic']:
                    text = f'*{text}*'
                line_text += text + ' '
            
            if line_text.strip():
                paragraph_lines.append(line_text.strip())
        
        paragraph_text = ' '.join(paragraph_lines)
        
        # Add HTML alignment for non-left aligned paragraphs in Markdown
        if alignment == 'center':
            return [f'<div align="center">{paragraph_text}</div>']
        elif alignment == 'right':
            return [f'<div align="right">{paragraph_text}</div>']
        
        return [paragraph_text]
    
    def _table_to_markdown(self, table_data):
        if not table_data or len(table_data) < 1:
            return ''
        
        markdown_lines = []
        
        # Header row
        if table_data:
            header = '| ' + ' | '.join([str(cell or '') for cell in table_data[0]]) + ' |'
            markdown_lines.append(header)
            
            # Separator
            separator = '|' + '---|' * len(table_data[0])
            markdown_lines.append(separator)
            
            # Data rows
            for row in table_data[1:]:
                row_text = '| ' + ' | '.join([str(cell or '') for cell in row]) + ' |'
                markdown_lines.append(row_text)
        
        return '\n'.join(markdown_lines)
    
    def to_docx(self):
        document = Document()
        
        if self.layout_mode == 'clean':
            return self._to_docx_clean(document)
        
        for block in self.formatted_content:
            if block.get('type') == 'table':
                self._add_table_to_docx(document, block['data'])
            else:
                block_content = block.get('content', [])
                block_type = block.get('type', 'paragraph')
                
                if block_type == 'heading':
                    self._add_heading_to_docx(document, block_content)
                elif block_type == 'list':
                    self._add_list_to_docx(document, block_content)
                else:
                    self._add_paragraph_to_docx(document, block_content)
        
        return document
    
    def _to_docx_clean(self, document):
        all_text = []
        for block in self.formatted_content:
            if block.get('type') == 'table':
                continue
            
            block_content = block.get('content', [])
            for line in block_content:
                line_text = ' '.join([span['text'] for span in line])
                if line_text.strip():
                    all_text.append(line_text.strip())
        
        paragraph = document.add_paragraph(' '.join(all_text))
        return document
    
    def _add_heading_to_docx(self, document, block_content):
        if not block_content or not block_content[0]:
            return
        
        heading_level = min(9, block_content[0][0].get('heading', 3))
        alignment = block_content[0][0].get('alignment', 'left')
        text_content = ' '.join([' '.join([span['text'] for span in line]) for line in block_content])
        
        heading = document.add_heading(text_content, level=heading_level)
        
        # Set alignment
        if alignment == 'center':
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif alignment == 'left':
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def _add_list_to_docx(self, document, block_content):
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Inches
        
        for line in block_content:
            if not line:
                continue
                
            first_span = line[0]
            if first_span.get('is_list'):
                # Remove bullet/number from text
                text_content = first_span['text']
                for bullet in self.bullet_patterns:
                    if text_content.startswith(bullet):
                        text_content = text_content[len(bullet):].strip()
                        break
                
                import re
                text_content = re.sub(r'^[\d\w]+[\.\)]\s*', '', text_content)
                
                # Create list paragraph
                paragraph = document.add_paragraph()
                paragraph.style = 'List Bullet'
                paragraph.paragraph_format.left_indent = Inches(0.25)
                
                # Add first span text
                run = paragraph.add_run(text_content)
                run.bold = first_span.get('bold', False)
                run.italic = first_span.get('italic', False)
                
                # Add remaining spans
                for span in line[1:]:
                    run = paragraph.add_run(' ' + span['text'])
                    run.bold = span.get('bold', False)
                    run.italic = span.get('italic', False)
                    run.font.size = Pt(span.get('size', 11))
            else:
                # Regular paragraph in list context
                self._add_paragraph_to_docx(document, [line])
    
    def _add_paragraph_to_docx(self, document, block_content):
        if not block_content:
            return
        
        # Get alignment from first line
        alignment = block_content[0][0].get('alignment', 'left') if block_content[0] else 'left'
        
        paragraph = document.add_paragraph()
        
        # Set alignment
        if alignment == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif alignment == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        for line in block_content:
            for span in line:
                run = paragraph.add_run(span['text'] + ' ')
                run.bold = span.get('bold', False)
                run.italic = span.get('italic', False)
                run.font.size = Pt(span.get('size', 11) if span.get('size', 0) > 8 else 11)
    
    def _add_table_to_docx(self, document, table_data):
        if not table_data:
            return
        
        # Create table
        table = document.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = str(cell_data or '')
                    
                    # Make first row bold (headers)
                    if i == 0:
                        for paragraph in row.cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

@app.route("/", methods=["GET"])
def index():
    return open("index.html", encoding="utf-8").read()

@app.route("/login", methods=["POST"])
def login():
    """Simulate user login"""
    username = request.json.get("username", "").strip()
    password = request.json.get("password", "").strip()
    
    # Simple authentication (in production, use proper auth)
    if username and password:
        session['user_id'] = username
        session['username'] = username
        return jsonify({
            "success": True,
            "message": "Logged in successfully",
            "username": username,
            "daily_limit": LOGGED_IN_DAILY_LIMIT
        })
    else:
        return jsonify({"success": False, "message": "Invalid credentials"}), 400

@app.route("/logout", methods=["POST"])
def logout():
    """Logout user"""
    session.pop('user_id', None)
    session.pop('username', None)
    return jsonify({
        "success": True,
        "message": "Logged out successfully",
        "daily_limit": GUEST_DAILY_LIMIT
    })

@app.route("/status", methods=["GET"])
def user_status():
    """Get current user status and export limits"""
    user_key = get_user_key()
    current_exports, limit = get_export_status(user_key)
    
    return jsonify({
        "logged_in": is_logged_in(),
        "username": session.get('username', ''),
        "daily_limit": limit,
        "current_exports": current_exports,
        "remaining_exports": limit - current_exports,
        "user_type": "Logged-in User" if is_logged_in() else "Guest User"
    })

@app.route("/viewer", methods=["GET"])
def viewer():
    return open("viewer.html", encoding="utf-8").read()

@app.route("/batch", methods=["GET"])
def batch():
    return open("batch.html", encoding="utf-8").read()

@app.route("/batch-convert", methods=["POST"])
def batch_convert():
    """Batch conversion endpoint for multiple PDF files"""
    try:
        pdf_files = request.files.getlist("pdf_files")
        formats_str = request.form.get("formats", '["txt"]')
        layout_mode = request.form.get("layout_mode", "preserve")
        
        # Check export limits for batch processing
        user_key = get_user_key()
        file_count = len([f for f in pdf_files if f.filename])
        
        # Check if user has enough exports remaining
        current_exports, limit = get_export_status(user_key)
        if current_exports + file_count > limit:
            remaining = limit - current_exports
            user_type = "Logged-in User" if is_logged_in() else "Guest User"
            
            if remaining <= 0:
                if is_logged_in():
                    message = f"Daily export limit reached ({limit} exports/day). Try again tomorrow."
                else:
                    message = f"You've reached your {limit} export limit. Log in to get 3x more exports per day!"
            else:
                message = f"Not enough exports remaining. You have {remaining} exports left but trying to process {file_count} files. Please reduce the number of files or try again tomorrow."
            
            return jsonify({
                "error": message,
                "limit_reached": True,
                "user_type": user_type,
                "current_count": current_exports,
                "limit": limit,
                "remaining": remaining,
                "requested": file_count
            }), 429
        
        print(f"DEBUG: Received {len(pdf_files)} files")
        print(f"DEBUG: Formats string: {formats_str}")
        print(f"DEBUG: Layout mode: {layout_mode}")
        
        # Parse formats safely
        try:
            formats = json.loads(formats_str)
        except:
            formats = ["txt"]
        
        if not pdf_files or len(pdf_files) == 0:
            return jsonify({"error": "No files provided"}), 400
        
        # Create a persistent temporary directory
        temp_dir = tempfile.mkdtemp()
        print(f"DEBUG: Using temp directory: {temp_dir}")
        
        try:
            converted_files = []
            
            # Process each PDF file
            for i, pdf_file in enumerate(pdf_files):
                if not pdf_file or pdf_file.filename == '':
                    print(f"DEBUG: Skipping empty file at index {i}")
                    continue
                
                print(f"DEBUG: Processing file {i+1}: {pdf_file.filename}")
                
                # Save uploaded file
                original_name = os.path.splitext(pdf_file.filename)[0]
                pdf_path = os.path.join(temp_dir, f"{original_name}.pdf")
                pdf_file.save(pdf_path)
                print(f"DEBUG: Saved PDF to {pdf_path}")
                
                try:
                    # Initialize formatter
                    formatter = PDFFormatter(pdf_path, layout_mode)
                    formatter.extract_with_formatting()
                    print(f"DEBUG: Extracted formatting for {pdf_file.filename}")
                    
                    # Convert to each requested format
                    for format_type in formats:
                        print(f"DEBUG: Converting {pdf_file.filename} to {format_type}")
                        
                        output_filename = f"{original_name}.{format_type}"
                        format_dir = os.path.join(temp_dir, format_type)
                        os.makedirs(format_dir, exist_ok=True)
                        output_path = os.path.join(format_dir, output_filename)
                        
                        if format_type == "txt":
                            # Generate plain text
                            if layout_mode == "preserve":
                                all_text = []
                                for block in formatter.formatted_content:
                                    if block.get('type') == 'table':
                                        continue
                                    
                                    block_content = block.get('content', [])
                                    for line in block_content:
                                        line_text = ' '.join([span['text'] for span in line])
                                        if line_text.strip():
                                            all_text.append(line_text.strip())
                                
                                txt_content = '\n\n'.join(all_text)
                            else:
                                # Simple text extraction
                                doc = fitz.open(pdf_path)
                                txt_content = ""
                                try:
                                    for page in doc:
                                        txt_content += page.get_text("text") + "\n"
                                finally:
                                    doc.close()
                            
                            with open(output_path, "w", encoding="utf-8") as f:
                                f.write(txt_content)
                        
                        elif format_type == "html":
                            html_content = formatter.to_html()
                            with open(output_path, "w", encoding="utf-8") as f:
                                f.write(html_content)
                        
                        elif format_type == "markdown":
                            markdown_content = formatter.to_markdown()
                            with open(output_path, "w", encoding="utf-8") as f:
                                f.write(markdown_content)
                        
                        elif format_type == "docx":
                            docx_document = formatter.to_docx()
                            docx_document.save(output_path)
                        
                        converted_files.append((format_type, output_path, output_filename))
                        print(f"DEBUG: Successfully converted to {format_type}")
                
                except Exception as file_error:
                    print(f"ERROR processing {pdf_file.filename}: {str(file_error)}")
                    # Continue with other files even if one fails
                    continue
            
            if not converted_files:
                return jsonify({"error": "No files could be processed"}), 400
            
            # Create ZIP file with all converted files
            zip_filename = f"batch_conversion_{int(time.time())}.zip"
            zip_path = os.path.join(temp_dir, zip_filename)
            print(f"DEBUG: Creating ZIP at {zip_path}")
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for format_type, file_path, filename in converted_files:
                    # Add file to ZIP with format-based folder structure
                    arcname = os.path.join(format_type, filename)
                    zipf.write(file_path, arcname)
                    print(f"DEBUG: Added {arcname} to ZIP")
                
                # Add a README file
                readme_content = f"""PDF2Preserve Batch Conversion Results
==========================================

Conversion Details:
- Files processed: {len([f for f in pdf_files if f.filename])}
- Formats: {', '.join(formats)}
- Layout mode: {layout_mode}
- Conversion time: {time.strftime('%Y-%m-%d %H:%M:%S')}

Folder Structure:
{chr(10).join([f'- {format_type}/ - Contains all .{format_type} files' for format_type in formats])}

Visit PDF2Preserve for more conversions!
"""
                zipf.writestr("README.txt", readme_content)
            
            print(f"DEBUG: ZIP file created successfully")
            
            # Update export count for processed files
            processed_file_count = len(set([filename.split('.')[0] for _, _, filename in converted_files]))
            for _ in range(processed_file_count):
                check_and_update_exports(user_key)
            
            # Return ZIP file
            def remove_temp_dir():
                try:
                    import shutil
                    shutil.rmtree(temp_dir)
                except:
                    pass
            
            # Schedule cleanup after response
            import atexit
            atexit.register(remove_temp_dir)
            
            return send_file(zip_path, as_attachment=True, download_name=zip_filename)
        
        except Exception as inner_error:
            print(f"INNER ERROR: {str(inner_error)}")
            # Clean up temp directory on error
            try:
                import shutil
                shutil.rmtree(temp_dir)
            except:
                pass
            raise inner_error
    
    except Exception as e:
        print(f"BATCH CONVERT ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Batch conversion failed: {str(e)}"}), 500

@app.route("/extract-text", methods=["POST"])
def extract_text():
    """Real-time text extraction endpoint for the side-by-side viewer"""
    try:
        file = request.files["pdf_file"]
        layout_mode = request.form.get("layout_mode", "preserve")
        
        if not file or file.filename == '':
            return jsonify({"error": "No file provided"}), 400
        
        # Save uploaded file temporarily
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)
        
        # Initialize formatter with layout mode
        formatter = PDFFormatter(filepath, layout_mode)
        formatter.extract_with_formatting()
        
        # Generate text in different formats
        txt_content = ""
        if layout_mode == "preserve":
            all_text = []
            for block in formatter.formatted_content:
                if block.get('type') == 'table':
                    continue  # Skip tables in text mode
                
                block_content = block.get('content', [])
                for line in block_content:
                    line_text = ' '.join([span['text'] for span in line])
                    if line_text.strip():
                        all_text.append(line_text.strip())
            
            txt_content = '\n\n'.join(all_text)
        else:
            # Simple text extraction
            doc = fitz.open(filepath)
            try:
                for page in doc:
                    txt_content += page.get_text("text") + "\n"
            finally:
                doc.close()
        
        html_content = formatter.to_html()
        markdown_content = formatter.to_markdown()
        
        # Clean up uploaded file
        try:
            time.sleep(0.1)
            os.remove(filepath)
        except PermissionError:
            time.sleep(0.5)
            try:
                os.remove(filepath)
            except PermissionError:
                pass  # Leave for manual cleanup
        
        return jsonify({
            "txt": txt_content,
            "html": html_content,
            "markdown": markdown_content,
            "success": True
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/upload", methods=["POST"])
def upload_pdf():
    # Check export limits before processing
    user_key = get_user_key()
    can_export, current_count, limit = check_and_update_exports(user_key)
    
    if not can_export:
        user_type = "Logged-in User" if is_logged_in() else "Guest User"
        if is_logged_in():
            message = f"Daily export limit reached ({limit} exports/day). Try again tomorrow."
        else:
            message = f"You've reached your {limit} export limit. Log in to get 3x more exports per day!"
        
        return jsonify({
            "error": message,
            "limit_reached": True,
            "user_type": user_type,
            "current_count": current_count,
            "limit": limit
        }), 429
    
    file = request.files["pdf_file"]
    export_format = request.form.get("export_format", "txt")
    layout_mode = request.form.get("layout_mode", "preserve")
    
    # Save uploaded file
    filepath = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(filepath)
    
    # Initialize formatter with layout mode
    formatter = PDFFormatter(filepath, layout_mode)
    
    if export_format == "txt":
        if layout_mode == "preserve":
            # Extract with structure and convert to plain text
            formatter.extract_with_formatting()
            all_text = []
            for block in formatter.formatted_content:
                if block.get('type') == 'table':
                    continue  # Skip tables in text mode
                
                block_content = block.get('content', [])
                for line in block_content:
                    line_text = ' '.join([span['text'] for span in line])
                    if line_text.strip():
                        all_text.append(line_text.strip())
            
            full_text = '\n\n'.join(all_text)
        else:
            # Simple text extraction (original behavior)
            doc = fitz.open(filepath)
            full_text = ""
            try:
                for page in doc:
                    full_text += page.get_text("text") + "\n"
            finally:
                doc.close()
        
        output_path = filepath.replace(".pdf", ".txt")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_text)
            
    elif export_format == "html":
        formatter.extract_with_formatting()
        html_content = formatter.to_html()
        
        output_path = filepath.replace(".pdf", ".html")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
            
    elif export_format == "markdown":
        formatter.extract_with_formatting()
        markdown_content = formatter.to_markdown()
        
        output_path = filepath.replace(".pdf", ".md")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
            
    elif export_format == "docx":
        formatter.extract_with_formatting()
        docx_document = formatter.to_docx()
        
        output_path = filepath.replace(".pdf", ".docx")
        docx_document.save(output_path)
    
    # Clean up uploaded PDF with retry mechanism
    try:
        time.sleep(0.1)  # Small delay to ensure file handles are released
        os.remove(filepath)
    except PermissionError:
        # If file is still locked, try again after a short delay
        time.sleep(0.5)
        try:
            os.remove(filepath)
        except PermissionError:
            # If still locked, leave it for manual cleanup
            pass
    
    return send_file(output_path, as_attachment=True)

if __name__ == "__main__":
    # For Railway deployment - use environment PORT or default to 5000
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
